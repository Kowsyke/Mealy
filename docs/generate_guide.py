#!/usr/bin/env python3
"""
Generate Mealy Project Guide (docx).
Output: docs/mealy_project_guide.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), "mealy_project_guide.docx")

doc = Document()

# ── page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── colour palette ────────────────────────────────────────────────────────────
ORANGE  = RGBColor(0xFF, 0x8C, 0x42)   # accent
DARK    = RGBColor(0x1A, 0x1A, 0x2E)   # headings
MID     = RGBColor(0x44, 0x44, 0x55)   # sub-headings
CODE_BG = RGBColor(0xF4, 0xF4, 0xF8)   # inline code bg

# ── helper functions ──────────────────────────────────────────────────────────
def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = DARK if level == 1 else MID
    if level == 1:
        run.font.size = Pt(20)
        run.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.bold = True
    else:
        run.font.size = Pt(12)
        run.bold = True
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after  = Pt(6)
    return p

def add_para(text, bold=False, italic=False, color=None, size=11, indent=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if indent:
        p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(3)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x55)
    # shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F8')
    pPr.append(shd)
    return p

def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    run = p.add_run('📌  ' + text)
    run.font.size  = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x88)
    p.paragraph_format.space_after = Pt(6)

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run('─' * 72)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xDD)
    run.font.size = Pt(8)

def add_screenshot_item(num, description, where, what):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'  Screenshot {num}: ')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = ORANGE
    p.add_run(description).font.size = Pt(11)
    doc.add_paragraph(f'     Where: {where}', style='List Bullet').runs[0].font.size = Pt(10)
    doc.add_paragraph(f'     What to show: {what}', style='List Bullet').runs[0].font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MEALY')
r.font.size = Pt(42)
r.font.bold = True
r.font.color.rgb = ORANGE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Food Recognition AI — Project Guide')
r.font.size = Pt(18)
r.font.color.rgb = DARK

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('CMS22202 · Operating Systems & Cloud Computing\nRavensbourne University London')
r.font.size = Pt(12)
r.font.color.rgb = MID

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('A complete step-by-step account of what was built,\nhow it works, and the concepts behind it.')
r.font.size = Pt(11)
r.font.italic = True
r.font.color.rgb = MID

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS  (manual — Word can update it)
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Contents', 1)
toc_items = [
    ('1', 'What is Mealy?'),
    ('2', 'Project Setup & Environment'),
    ('3', 'Understanding the Datasets'),
    ('4', 'Building the Data Pipeline'),
    ('5', 'Designing the Neural Network Model'),
    ('6', 'Training the Model'),
    ('7', 'The Flask REST API'),
    ('8', 'The Web Interface (Frontend)'),
    ('9', 'Deployment Architecture Overview'),
    ('10', 'SSH Reverse Tunnel — Exposing a Local Server'),
    ('11', 'Systemd Services — Auto-Start on Login'),
    ('12', 'Nginx Reverse Proxy on EC2'),
    ('13', 'HTTPS with Let\'s Encrypt & sslip.io'),
    ('14', 'Kowsite Integration'),
    ('15', 'Key Concepts Glossary'),
    ('16', 'Screenshots to Take for Your Report'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f'  Step {num}. ')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = ORANGE
    p.add_run(title).font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 1: WHAT IS MEALY?
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Step 1 — What is Mealy?', 1)

add_para(
    'Mealy is a real-time food recognition system built as the practical component of the '
    'CMS22202 Operating Systems & Cloud Computing module at Ravensbourne University London. '
    'It uses deep learning to identify food items through a device camera and estimate their '
    'caloric content — all running live in the browser.'
)

add_heading('What the system does', 2)
add_bullet('Activates your device camera through a web browser')
add_bullet('Captures frames when you press "Scan Frame" (or automatically every 2 seconds)')
add_bullet('Sends the image to a Python AI server running on a local Fedora Linux machine')
add_bullet('The AI identifies the food using a trained neural network')
add_bullet('Returns the food name, confidence percentage, and estimated calories')
add_bullet('Displays results as cards in real time in the browser')

add_heading('Technologies used', 2)
tech = [
    ('TensorFlow 2.21', 'Deep learning framework — trains and runs the neural network'),
    ('MobileNetV2', 'Pre-trained image recognition model used as the "brain base"'),
    ('Flask', 'Lightweight Python web server — serves the AI as a REST API'),
    ('Food-101', 'Dataset of 101,000 labelled food images (101 classes)'),
    ('Fruit-360', 'Dataset of 90,000 fruit/vegetable images (36 classes)'),
    ('HTML / CSS / JS', 'Browser frontend — camera, UI, results display'),
    ('SSH Reverse Tunnel', 'Connects the local Flask server to a public EC2 server'),
    ('Amazon EC2', 'Cloud server that makes the app accessible on the internet'),
    ('Nginx', 'Web server on EC2 — routes traffic and forwards to Flask'),
    ('Let\'s Encrypt', 'Free HTTPS certificate (required for camera access in browsers)'),
    ('systemd', 'Linux service manager — auto-starts Flask and the SSH tunnel'),
]
for name, desc in tech:
    add_bullet(f'{desc}', bold_prefix=f'{name}: ')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 2: PROJECT SETUP
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Step 2 — Project Setup & Environment', 1)

add_para(
    'Before any code could be written, the development environment had to be prepared. '
    'This is one of the most important (and underappreciated) steps in any software project.'
)

add_heading('Creating the project structure', 2)
add_para('The project lives at:')
add_code('/home/K/Storage/Projects/Mealy/')
add_para('Key folders and files:')
structure = [
    ('app.py', 'The Flask API — the main server'),
    ('model.py', 'Neural network definition (MobileNetV2 + custom head)'),
    ('train.py / train_optimized.py', 'Training scripts'),
    ('load_data.py', 'Loads and pipelines the dataset'),
    ('preprocess.py', 'Image resizing and normalisation'),
    ('detect.py', 'Prediction logic — runs inference on an image'),
    ('calories.py / fruit_calories.py', 'Calorie lookup tables'),
    ('ui/index.html', 'The web frontend (camera, scan button, results)'),
    ('food-101/', 'The Food-101 dataset (downloaded ~5 GB)'),
    ('Fruit_dataset/', 'The Fruit-360 dataset'),
    ('models/', 'Saved trained model weights (.keras files)'),
    ('keggle/', 'Optimised retrain outputs'),
]
for fname, desc in structure:
    add_bullet(desc, bold_prefix=f'{fname} — ')

add_heading('Python virtual environment', 2)
add_para(
    'A virtual environment isolates your project\'s dependencies from the rest of the system, '
    'so different projects can use different library versions without conflict.'
)
add_code('python3 -m venv venv_archive')
add_code('source venv_archive/bin/activate')
add_code('pip install tensorflow flask flask-cors pillow numpy')
add_note('TensorFlow 2.21.0 was used. TF versions can break APIs — always pin your version.')

add_heading('Why no GPU?', 2)
add_para(
    'The Fedora development machine has no dedicated GPU, so TensorFlow runs on CPU. '
    'Training on CPU for 50 epochs over 75,000 images would take many hours. '
    'Google Colab (free cloud GPUs) was used for training — the trained model weights '
    'were then downloaded and used locally for inference (running predictions).'
)
add_note('Inference (predicting from a single image) is fast even on CPU — typically under 1 second.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 3: DATASETS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Step 3 — Understanding the Datasets', 1)

add_heading('Food-101', 2)
add_para(
    'Food-101 is a benchmark dataset released by ETH Zürich. It contains exactly 1,000 '
    'images per class across 101 food categories.'
)
add_bullet('101 food classes (pizza, sushi, steak, ice cream, etc.)')
add_bullet('75,750 training images')
add_bullet('25,250 test images')
add_bullet('All images are JPEG, varied sizes (resized to 224×224 for training)')
add_bullet('Downloaded from the official source as a .tar.gz archive (~5 GB)')

add_para('The dataset folder structure:')
add_code('food-101/')
add_code('  images/  ← one subfolder per class, e.g. images/pizza/001.jpg')
add_code('  meta/')
add_code('    classes.txt    ← 101 class names, one per line')
add_code('    train.txt      ← list of image paths for training')
add_code('    test.txt       ← list of image paths for testing')

add_heading('Fruit-360 (subset)', 2)
add_para(
    'A Kaggle dataset of fruit and vegetable images taken against white backgrounds. '
    'A curated subset of 36 classes was used to complement the Food-101 model.'
)
add_bullet('36 classes (apple, banana, mango, tomato, broccoli, etc.)')
add_bullet('~90 images per class at 100×100 pixels')
add_bullet('Very high accuracy (96.7%) due to consistent backgrounds')

add_heading('Why two datasets?', 2)
add_para(
    'Food-101 covers cooked/prepared foods (dishes, meals). Fruit-360 covers raw produce. '
    'Using both gives broader coverage. The system supports three modes:'
)
add_bullet('Food mode — 101 classes, targets prepared meals')
add_bullet('Fruit mode — 36 classes, targets raw fruit/veg')
add_bullet('Combined mode — both models, returns the highest-confidence result')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 4: DATA PIPELINE
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Step 4 — Building the Data Pipeline', 1)

add_para(
    'A data pipeline efficiently loads, preprocesses, and feeds images to the model during '
    'training. TensorFlow\'s tf.data API is used — it runs preprocessing in parallel on multiple '
    'CPU cores and keeps the GPU (or CPU-based training) constantly fed without waiting for disk I/O.'
)

add_heading('preprocess.py — image loading and normalisation', 2)
add_para('Every image goes through these steps:')
add_bullet('Read the JPEG file from disk')
add_bullet('Decode it into a 3-channel (RGB) pixel array')
add_bullet('Resize it to 224 × 224 pixels (required by MobileNetV2)')
add_bullet('Normalise pixel values from [0, 255] to [0.0, 1.0] (divide by 255)')
add_code('# preprocess.py — key function')
add_code('def preprocess_path(path):')
add_code('    raw = tf.io.read_file(path)')
add_code('    image = tf.image.decode_jpeg(raw, channels=3)')
add_code('    image = tf.image.resize(image, [224, 224])')
add_code('    return tf.cast(image, tf.float32) / 255.0')
add_note(
    'IMPORTANT: .numpy() must NOT be called here. This function runs inside '
    'tf.data.Dataset.map() which operates in TensorFlow graph mode — '
    'everything must stay as symbolic tensors.'
)

add_heading('load_data.py — building the dataset pipeline', 2)
add_para('The pipeline for training:')
add_bullet('Read train.txt / test.txt to get image paths and class labels')
add_bullet('Create a tf.data.Dataset from the list of file paths')
add_bullet('Apply preprocess_path() to each path in parallel (num_parallel_calls=AUTOTUNE)')
add_bullet('Zip images with their labels')
add_bullet('Apply augmentation to training images (not test images)')
add_bullet('Shuffle (training only) → batch into groups of 32 → prefetch')

add_heading('Data augmentation', 2)
add_para(
    'Augmentation artificially increases dataset diversity by applying random transformations '
    'to training images. This prevents overfitting (memorising the training data) and helps '
    'the model generalise to new images it has never seen.'
)
add_bullet('Random horizontal flip — the model sees mirrored versions of food')
add_bullet('Random brightness adjustment (±15%) — simulates different lighting')
add_bullet('Random contrast adjustment (±15%) — simulates camera variations')
add_bullet('Random saturation adjustment — colour variation')
add_bullet('Random crop to 90% then resize back — simulates zoom variation')
add_code('# Augmentation uses only stateless tf.image ops — no Keras layers')
add_code('# (Keras layers create tf.Variables inside tf.function, which is forbidden)')
add_code('image = tf.image.random_flip_left_right(image)')
add_code('image = tf.image.random_crop(image, [int(224*0.9), int(224*0.9), 3])')
add_code('image = tf.image.resize(image, [224, 224])')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 5: MODEL
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Step 5 — Designing the Neural Network Model', 1)

add_heading('What is transfer learning?', 2)
add_para(
    'Training a deep neural network from scratch requires millions of images and weeks of '
    'GPU time. Transfer learning is the technique of taking a model already trained on a '
    'huge dataset and adapting it to a new task.'
)
add_para(
    'Think of it like hiring an experienced surgeon to learn a new procedure rather than '
    'training someone from medical school. The surgeon already understands anatomy '
    '(low-level features like edges, textures, shapes) — they only need to learn the '
    'new specific skill.'
)

add_heading('MobileNetV2 — the base model', 2)
add_para(
    'MobileNetV2 is a convolutional neural network designed by Google. It was pre-trained '
    'on ImageNet — a dataset of 1.2 million images across 1,000 categories. It is efficient '
    'enough to run on mobile devices (hence "Mobile") while being highly accurate.'
)
add_bullet('Input: 224 × 224 × 3 image (width × height × RGB channels)')
add_bullet('Architecture: series of depthwise-separable convolution layers')
add_bullet('Output: 1,280-dimensional feature vector (a compressed description of the image)')
add_bullet('Pre-trained weights capture shapes, textures, colours, patterns from ImageNet')

add_heading('Custom classification head', 2)
add_para('On top of MobileNetV2, a custom head was added:')
add_bullet('GlobalAveragePooling2D — flattens the 7×7×1280 feature maps to 1280 numbers')
add_bullet('Dense(256, activation=relu) — learns food-specific combinations of features')
add_bullet('Dropout(0.3) — randomly drops 30% of neurons during training to prevent overfitting')
add_bullet('Dense(101, activation=softmax) — outputs a probability for each of the 101 food classes')
add_code('# model.py — core architecture')
add_code('base = MobileNetV2(input_shape=(224,224,3), include_top=False, weights="imagenet")')
add_code('x = GlobalAveragePooling2D()(base.output)')
add_code('x = Dense(256, activation="relu")(x)')
add_code('x = Dropout(0.3)(x)')
add_code('out = Dense(101, activation="softmax")(x)')
add_code('model = Model(inputs=base.input, outputs=out)')

add_heading('What is softmax?', 2)
add_para(
    'Softmax converts the final layer\'s raw numbers into probabilities that all add up to 1.0. '
    'If the model outputs [0.72, 0.15, 0.08, ...] for [pizza, burger, salad, ...], it is '
    '72% confident the image shows pizza.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 6: TRAINING
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Step 6 — Training the Model', 1)

add_para(
    'Training is the process where the model learns by looking at labelled examples and '
    'adjusting its internal parameters (weights) to reduce its prediction errors.'
)

add_heading('Two-phase training strategy', 2)
add_para(
    'Training was split into two phases. This is a standard transfer learning approach:'
)

add_heading('Phase 1 — Frozen base (20 epochs)', 3)
add_bullet('The MobileNetV2 base layers are frozen (their weights do not change)')
add_bullet('Only the custom head (Dense + Dropout + output layer) is trained')
add_bullet('Learning rate: 0.001 (relatively fast learning)')
add_bullet('Reason: the base already knows good image features; we just teach it to classify food')
add_note('An epoch is one complete pass through all 75,750 training images.')

add_heading('Phase 2 — Fine-tuning (30 epochs)', 3)
add_bullet('The top 30 layers of MobileNetV2 are unfrozen')
add_bullet('Both the custom head AND the top base layers are trained together')
add_bullet('Learning rate: 0.00001 (10× slower — to avoid overwriting learned features)')
add_bullet('Reason: the model now fine-tunes the high-level feature detectors for food specifically')

add_heading('Training callbacks (automatic helpers)', 2)
callbacks = [
    ('EarlyStopping', 'Monitors validation accuracy. Stops training if it hasn\'t improved for 7 epochs. Saves the best weights.'),
    ('ReduceLROnPlateau', 'Halves the learning rate if accuracy stagnates for 4 epochs. Helps the model break out of flat spots.'),
    ('ModelCheckpoint', 'Saves the model weights after every epoch, so training can be resumed if interrupted.'),
]
for name, desc in callbacks:
    add_bullet(desc, bold_prefix=f'{name}: ')

add_heading('Running on Google Colab', 2)
add_para(
    'Colab provides free access to NVIDIA T4/A100 GPUs. Training that would take 8+ hours '
    'on CPU completes in 30–60 minutes on a Colab GPU.'
)
add_bullet('Upload the dataset to Google Drive or download it directly in Colab')
add_bullet('Run train_optimized.py in a Colab notebook cell')
add_bullet('Download the .keras model file when training finishes')
add_bullet('Place the model in the models/ or keggle/ folder on the local machine')

add_heading('Results achieved', 2)
add_bullet('Food model (MobileNetV2, Food-101): ~37.6% top-1 accuracy on test set')
add_bullet('Fruit model (fine-tuned, Fruit-360 subset): ~96.7% top-1 accuracy')
add_note(
    '37.6% for 101 classes is reasonable — random chance would be 0.99%. '
    'State-of-the-art models reach ~90% but require 10× more compute and data.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 7: FLASK API
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Step 7 — The Flask REST API', 1)

add_para(
    'Flask is a lightweight Python web framework. It turns the trained model into an HTTP '
    'service that any browser or app can talk to by sending HTTP requests.'
)

add_heading('What is a REST API?', 2)
add_para(
    'REST (Representational State Transfer) is a style for building web services. '
    'The client (browser) sends an HTTP request to a URL (called an endpoint) with data '
    'attached, and the server responds with JSON.'
)

add_heading('API endpoints', 2)
endpoints = [
    ('GET  /health', 'Returns JSON showing API status and which models are loaded'),
    ('GET  /', 'Serves ui/index.html — the web frontend'),
    ('POST /detect', 'Upload an image → returns food detections with confidence and calories'),
    ('POST /detect_fruit', 'Same but uses the fruit/vegetable model'),
    ('POST /detect_combined', 'Runs both models and returns the combined best results'),
]
for ep, desc in endpoints:
    add_bullet(desc, bold_prefix=f'{ep}: ')

add_heading('Lazy model loading', 2)
add_para(
    'Loading a TensorFlow model from disk takes several seconds. Flask starts in under '
    '0.35 seconds by NOT loading models at startup. Instead, models load on the first '
    'request to /detect or /detect_fruit.'
)
add_code('# app.py — lazy loading pattern')
add_code('food_model = None  # not loaded yet')
add_code('')
add_code('@app.route("/detect", methods=["POST"])')
add_code('def detect():')
add_code('    global food_model')
add_code('    if food_model is None:')
add_code('        food_model = tf.keras.models.load_model("models/food_model.keras")')
add_code('    # ... run inference ...')

add_heading('How a scan request works', 2)
add_bullet('Browser captures a video frame and converts it to JPEG')
add_bullet('JPEG is sent as a multipart/form-data POST to /detect')
add_bullet('Flask receives the image, passes it through preprocess → model → softmax')
add_bullet('Top predictions above 15% confidence are returned as JSON')
add_bullet('Calorie estimates are looked up in a dictionary (calories.py)')
add_bullet('Browser renders the results as detection cards')

add_heading('CORS — allowing the browser to call the API', 2)
add_para(
    'Browsers block JavaScript from calling a different domain than the page was loaded from '
    '(same-origin policy). Flask-CORS adds the right HTTP headers to allow cross-origin requests.'
)
add_code('from flask_cors import CORS')
add_code('CORS(app)  # allow all origins')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 8: FRONTEND
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Step 8 — The Web Interface (Frontend)', 1)

add_para(
    'The frontend (ui/index.html) is a single HTML file containing all CSS and JavaScript. '
    'It runs entirely in the browser — no separate build step or framework needed.'
)

add_heading('Layout: three-column grid', 2)
add_bullet('Left panel (200px): System status, model selector, session stats')
add_bullet('Centre panel (flexible): Live camera feed, scan controls')
add_bullet('Right panel (280px): Detection results cards, total calorie count')
add_code('/* CSS Grid — the layout skeleton */')
add_code('.app {')
add_code('  display: grid;')
add_code('  grid-template-columns: 200px 1fr 280px;')
add_code('  height: 100vh;')
add_code('}')

add_heading('Camera access — getUserMedia()', 2)
add_para(
    'The browser API getUserMedia() requests permission to access the device camera. '
    'This is what shows the permission dialog ("Allow this site to use your camera?").'
)
add_code('const stream = await navigator.mediaDevices.getUserMedia({')
add_code('  video: { width: { ideal: 1280 }, height: { ideal: 720 } },')
add_code('  audio: false')
add_code('});')
add_code('video.srcObject = stream;  // display in <video> element')
add_note(
    'getUserMedia() ONLY works on HTTPS (or localhost). This is a browser security '
    'requirement — without HTTPS, the browser silently refuses without showing a dialog. '
    'This is why HTTPS (Let\'s Encrypt) had to be set up on the EC2 server.'
)

add_heading('Capturing and sending a frame', 2)
add_para('When "Scan Frame" is pressed:')
add_bullet('A hidden HTML <canvas> element is created')
add_bullet('The current video frame is drawn onto the canvas')
add_bullet('The canvas is converted to a JPEG blob (compressed image)')
add_bullet('The blob is sent as a POST request to /detect using the Fetch API')
add_bullet('The JSON response is parsed and rendered as detection cards')

add_heading('Camera toggle (System section)', 2)
add_para(
    'The camera status indicator in the System section is interactive. '
    'Clicking it toggles the camera on or off:'
)
add_bullet('Click when camera is off → calls getUserMedia() → browser asks for permission')
add_bullet('Click when camera is on → stops all camera tracks → camera light turns off')
add_bullet('Camera also auto-starts when the page first loads')

add_heading('API detection (how the URL is determined)', 2)
add_code('const API = (() => {')
add_code('  if (location.hostname === "localhost") return "http://localhost:5001";')
add_code('  if (location.pathname.startsWith("/mealy")) return location.origin + "/mealy";')
add_code('  return location.origin;')
add_code('})();')
add_para(
    'This automatically points to the right server whether you\'re running locally, '
    'accessing via the EC2 IP, or via the sslip.io HTTPS domain.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 9: DEPLOYMENT ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Step 9 — Deployment Architecture Overview', 1)

add_para('The complete data flow when a user scans food through Kowsite:')

add_code('User\'s browser  →  https://56-228-34-22.sslip.io/mealy/')
add_code('                       ↓')
add_code('          EC2 nginx (port 443, HTTPS)')
add_code('                       ↓')
add_code('          nginx proxies /mealy/ → localhost:5001')
add_code('                       ↓')
add_code('          SSH reverse tunnel (port 5001 on EC2')
add_code('          is forwarded to port 5001 on Fedora)')
add_code('                       ↓')
add_code('          Flask API on Fedora (port 5001)')
add_code('                       ↓')
add_code('          TensorFlow model runs inference')
add_code('                       ↓')
add_code('          JSON response travels back the same path')

add_para(
    'This architecture means the AI runs on your powerful local Fedora machine (with all '
    'its RAM and storage) but is accessible anywhere in the world through the EC2 cloud server.'
)

add_heading('Why this architecture?', 2)
add_bullet('EC2 has a public IP — your home/university machine does not')
add_bullet('TensorFlow models are large — keeping them on a local machine avoids cloud storage costs')
add_bullet('The SSH tunnel is encrypted end-to-end (no sensitive data travels unprotected)')
add_bullet('If the local machine is off, the tunnel drops gracefully and EC2 shows an error')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 10: SSH REVERSE TUNNEL
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Step 10 — SSH Reverse Tunnel', 1)

add_heading('What is SSH?', 2)
add_para(
    'SSH (Secure Shell) is a protocol for securely connecting to remote computers over a '
    'network. It is encrypted and authenticated using key pairs.'
)

add_heading('What is port forwarding?', 2)
add_para(
    'Normally SSH lets you run commands on a remote machine. But it can also create a '
    '"tunnel" — a pipe that forwards network traffic from one machine to another through '
    'the encrypted SSH connection.'
)

add_heading('Reverse tunnel explained', 2)
add_para(
    'A reverse tunnel makes a port on the remote machine (EC2) forward to a port on '
    'the local machine (Fedora). This is the opposite of a normal tunnel.'
)
add_code('ssh -N -R 5001:localhost:5001 ubuntu@56.228.34.22 -i ~/Keys/fedora-access.pem')
add_para('Breaking down the command:')
add_bullet('-N: do not execute a remote command (just keep the tunnel open)')
add_bullet('-R 5001:localhost:5001: on EC2, port 5001 → Fedora port 5001')
add_bullet('ubuntu@56.228.34.22: connect to EC2 as ubuntu user')
add_bullet('-i ~/Keys/fedora-access.pem: use this private key for authentication (no password)')

add_note(
    'The result: anything that connects to localhost:5001 on EC2 gets transparently '
    'forwarded to localhost:5001 on your Fedora machine — as if Flask were running on EC2.'
)

add_heading('Keeping the tunnel alive', 2)
add_bullet('ServerAliveInterval=30: send a keepalive packet every 30 seconds')
add_bullet('ServerAliveCountMax=3: give up after 3 missed keepalives (90 seconds)')
add_bullet('ExitOnForwardFailure=yes: fail immediately if the port is already in use on EC2')
add_bullet('Restart=always in systemd: if the SSH process dies, systemd restarts it after 10 seconds')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 11: SYSTEMD
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Step 11 — Systemd Services', 1)

add_para(
    'Systemd is the Linux system and service manager. It starts, stops, and monitors '
    'processes. User services (systemd --user) run under your own account without needing root.'
)

add_heading('mealy-api.service', 2)
add_para(
    'Starts the Flask API automatically when you log into your Fedora session.'
)
add_code('[Unit]')
add_code('Description=Mealy Food Recognition API')
add_code('After=network.target')
add_code('')
add_code('[Service]')
add_code('Type=simple')
add_code('WorkingDirectory=/home/K/Storage/Projects/Mealy')
add_code('ExecStart=/path/to/venv/bin/python3 app.py')
add_code('Restart=on-failure')
add_code('')
add_code('[Install]')
add_code('WantedBy=default.target')

add_heading('mealy-tunnel.service', 2)
add_para(
    'Starts the SSH reverse tunnel after mealy-api is confirmed running. '
    'The Requires= line means: if the API is not running, don\'t start the tunnel.'
)
add_code('[Unit]')
add_code('Requires=mealy-api.service    # tunnel only makes sense if API is up')
add_code('After=mealy-api.service network-online.target')
add_code('')
add_code('[Service]')
add_code('ExecStart=/usr/bin/ssh -N -R 5001:localhost:5001 \\')
add_code('    -o ServerAliveInterval=30 \\')
add_code('    -o ExitOnForwardFailure=yes \\')
add_code('    -i /home/K/Projects/Keys/keys/fedora-access.pem \\')
add_code('    ubuntu@56.228.34.22')
add_code('Restart=always')
add_code('RestartSec=10')

add_heading('Useful systemd commands', 2)
cmds = [
    ('systemctl --user status mealy-api', 'Check if the API is running'),
    ('systemctl --user status mealy-tunnel', 'Check if the tunnel is active'),
    ('systemctl --user restart mealy-api', 'Restart the Flask API'),
    ('journalctl --user -u mealy-api -f', 'Follow the API logs live'),
    ('systemctl --user enable mealy-api', 'Enable auto-start on login'),
]
for cmd, desc in cmds:
    add_bullet(desc, bold_prefix=f'{cmd}: ')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 12: NGINX
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Step 12 — Nginx Reverse Proxy on EC2', 1)

add_para(
    'Nginx is a high-performance web server and reverse proxy. On the EC2 instance, '
    'it listens on ports 80 (HTTP) and 443 (HTTPS) and routes incoming requests to '
    'the appropriate backend services.'
)

add_heading('What is a reverse proxy?', 2)
add_para(
    'A reverse proxy sits in front of your application and forwards requests to it. '
    'The client only talks to nginx — it never knows about Flask or port 5001. '
    'Benefits: security, SSL termination, routing, load balancing.'
)

add_heading('The /mealy/ location block', 2)
add_code('location /mealy/ {')
add_code('    proxy_pass         http://localhost:5001/;')
add_code('    proxy_http_version 1.1;')
add_code('    proxy_set_header   Host              $host;')
add_code('    proxy_set_header   X-Real-IP         $remote_addr;')
add_code('    client_max_body_size 20m;   # allow large image uploads')
add_code('    proxy_read_timeout   120s;  # wait up to 2 min for model inference')
add_code('}')

add_para('Path rewriting — important detail:')
add_bullet('Browser requests /mealy/detect')
add_bullet('Nginx strips the /mealy/ prefix and forwards /detect to Flask')
add_bullet('Flask receives /detect — exactly what it expects')
add_bullet('This is achieved by the trailing slash: proxy_pass http://localhost:5001/ (note the /)')

add_heading('The Kowsite (Node.js)', 2)
add_para(
    'Kowsite is a separate Node.js application running on port 3000. '
    'Nginx forwards all other traffic (location /) to it.'
)
add_code('location / {')
add_code('    proxy_pass http://localhost:3000;')
add_code('    proxy_set_header Upgrade $http_upgrade;')
add_code('    proxy_set_header Connection "Upgrade";')
add_code('}')
add_note('The Upgrade headers are for WebSocket support — Kowsite uses WebSockets for real-time node monitoring.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 13: HTTPS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Step 13 — HTTPS with Let\'s Encrypt & sslip.io', 1)

add_heading('Why HTTPS is mandatory for camera access', 2)
add_para(
    'Browsers enforce a security policy called "Secure Context". The getUserMedia() API '
    '(camera access) only works on:'
)
add_bullet('localhost (your own machine, always trusted)')
add_bullet('Any origin served over HTTPS with a valid certificate')
add_para(
    'Without HTTPS, the browser silently refuses the camera request — no dialog, '
    'no error message to the user. This is why the camera appeared broken before HTTPS was added.'
)

add_heading('The problem: getting HTTPS for an IP address', 2)
add_para(
    'Let\'s Encrypt issues certificates for domain names, not bare IP addresses. '
    'The EC2 instance has the IP 56.228.34.22 but no domain name.'
)

add_heading('The solution: sslip.io', 2)
add_para(
    'sslip.io is a free wildcard DNS service. It maps any IP address to a subdomain automatically.'
)
add_code('56-228-34-22.sslip.io  →  resolves to  →  56.228.34.22')
add_para(
    'No registration needed — the mapping is automatic. Now Let\'s Encrypt can verify '
    'that we control the domain (via HTTP challenge) and issue a certificate.'
)

add_heading('Certbot — obtaining the certificate', 2)
add_code('sudo apt install certbot python3-certbot-nginx')
add_code('sudo certbot --nginx -d 56-228-34-22.sslip.io --redirect')
add_para('What certbot does automatically:')
add_bullet('Proves to Let\'s Encrypt that you control the domain (HTTP-01 challenge)')
add_bullet('Downloads a signed certificate valid for 90 days')
add_bullet('Patches the nginx configuration to use SSL on port 443')
add_bullet('Adds an HTTP → HTTPS redirect for all requests')
add_bullet('Sets up a cron job / systemd timer to auto-renew before expiry')

add_heading('The nginx HTTPS configuration (auto-generated)', 2)
add_code('server {')
add_code('    listen 443 ssl;')
add_code('    server_name 56-228-34-22.sslip.io;')
add_code('    ssl_certificate     /etc/letsencrypt/live/56-228-34-22.sslip.io/fullchain.pem;')
add_code('    ssl_certificate_key /etc/letsencrypt/live/56-228-34-22.sslip.io/privkey.pem;')
add_code('    # ... location blocks ...')
add_code('}')
add_code('')
add_code('server {')
add_code('    listen 80;')
add_code('    return 301 https://56-228-34-22.sslip.io$request_uri;  # HTTP → HTTPS')
add_code('}')

add_note('The certificate expires after 90 days but certbot auto-renews it. No manual action needed.')

add_heading('AWS Security Group', 2)
add_para(
    'AWS Security Groups are virtual firewalls for EC2 instances. Port 443 had to be '
    'explicitly opened in the security group (launch-wizard-2) — '
    'without this, all HTTPS traffic was silently dropped at the AWS network level, '
    'even though nginx was ready to receive it.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 14: KOWSITE INTEGRATION
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Step 14 — Kowsite Integration', 1)

add_para(
    'Kowsite is a custom operating system dashboard (Node.js + WebSocket) also running '
    'on the EC2 instance. Mealy was integrated as a "client application" within Kowsite.'
)

add_heading('What was added to Kowsite', 2)
add_bullet('A "Clients" section in the sidebar with a Mealy node card (styled like the other EC2 agent cards)')
add_bullet('Clicking the MEALY card opens a full-screen panel overlay (below the 48px header)')
add_bullet('The panel contains an iframe that loads /mealy/ — the Mealy web UI')
add_bullet('Camera access works because the iframe has allow="camera;microphone" and the parent page is HTTPS')

add_heading('The Mealy card HTML', 2)
add_code('<div class="nc" id="tnMealy" onclick="showPanel(\'mealy\')">')
add_code('  <div class="nc-top">')
add_code('    <span class="nc-dot"></span>')
add_code('    <span class="nc-name">MEALY</span>')
add_code('    <span class="nc-load">food ai</span>')
add_code('  </div>')
add_code('  <div class="nc-foot">localhost:5001</div>')
add_code('</div>')

add_heading('The Mealy panel (iframe)', 2)
add_code('<div id="mealyPanel" style="position:fixed;top:48px;left:0;right:0;bottom:0;z-index:99">')
add_code('  <iframe')
add_code('    id="mealyFrame"')
add_code('    src="about:blank"')
add_code('    allow="camera;microphone"')
add_code('    style="width:100%;height:100%;border:none"')
add_code('  ></iframe>')
add_code('</div>')

add_heading('Lazy loading the iframe', 2)
add_para(
    'The iframe starts as "about:blank" (empty). The /mealy/ URL is only injected '
    'when the user actually opens the Mealy panel. This avoids loading Flask/TF on every '
    'Kowsite pageload.'
)
add_code('if (name === "mealy") {')
add_code('  const f = document.getElementById("mealyFrame");')
add_code('  if (f && f.src === "about:blank") f.src = "/mealy/";')
add_code('}')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 15: GLOSSARY
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Step 15 — Key Concepts Glossary', 1)

glossary = [
    ('API (Application Programming Interface)',
     'A defined way for software components to communicate. The Flask API '
     'lets the browser send images and receive JSON results.'),
    ('Convolutional Neural Network (CNN)',
     'A type of neural network designed for image data. It uses sliding filters '
     '(kernels) to detect patterns like edges, textures, and shapes at different '
     'scales across an image.'),
    ('Transfer Learning',
     'Reusing a model pre-trained on one task (ImageNet classification) for a '
     'related task (food recognition). Saves enormous amounts of training time and data.'),
    ('Epoch',
     'One complete pass through the entire training dataset. After each epoch the '
     'model\'s weights are updated based on all the errors it made.'),
    ('Overfitting',
     'When a model memorises the training data instead of learning general patterns. '
     'It performs well on training images but poorly on new images it hasn\'t seen.'),
    ('Dropout',
     'A regularisation technique that randomly deactivates a fraction of neurons '
     'during training. Forces the network to learn robust, distributed representations.'),
    ('Softmax',
     'A function that converts raw model outputs into a probability distribution. '
     'All outputs sum to 1.0. The highest probability is the model\'s prediction.'),
    ('REST',
     'A web architecture style where operations are performed by HTTP verbs (GET, POST, '
     'PUT, DELETE) on resource URLs. POST /detect sends an image and gets food results.'),
    ('SSH (Secure Shell)',
     'A cryptographic protocol for secure remote access and file transfer. '
     'Also supports port forwarding (tunnels).'),
    ('Reverse Tunnel',
     'A tunnel where the remote machine\'s port forwards to the local machine. '
     'Makes a local server accessible through a remote server.'),
    ('Nginx',
     'A high-performance web server used here as a reverse proxy — routes incoming '
     'HTTPS requests to the appropriate local service (Flask or Node.js).'),
    ('TLS / HTTPS',
     'Transport Layer Security encrypts web traffic. HTTPS = HTTP + TLS. '
     'Required for browser camera access (getUserMedia) on non-localhost origins.'),
    ('Let\'s Encrypt',
     'A free, automated Certificate Authority. Issues TLS certificates valid for 90 days, '
     'automatically renewed by Certbot.'),
    ('systemd',
     'Linux\'s standard init system and service manager. Controls when processes start, '
     'monitors them, and restarts them on failure.'),
    ('tf.data',
     'TensorFlow\'s dataset pipeline API. Enables efficient parallel data loading, '
     'preprocessing, augmentation, batching, and prefetching.'),
    ('CORS (Cross-Origin Resource Sharing)',
     'A browser security mechanism. Flask-CORS adds HTTP headers that allow the '
     'browser to call the API from a different origin (domain/port).'),
    ('getUserMedia()',
     'A browser JavaScript API for accessing the device camera or microphone. '
     'Returns a MediaStream object that can be displayed in a <video> element.'),
    ('iframe',
     'An HTML element that embeds another web page inside the current page. '
     'Mealy runs inside an iframe within Kowsite.'),
]
for term, definition in glossary:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(term + ': ')
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = DARK
    p.add_run(definition).font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  STEP 16: SCREENSHOTS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('Step 16 — Screenshots to Take for Your Report', 1)

add_para(
    'Below is a prioritised list of screenshots for the assessment report. '
    'Take them in order — some require the system to be running live.'
)

add_heading('A. System Architecture & Infrastructure', 2)

add_screenshot_item(1,
    'HTTPS Mealy UI in browser',
    'Navigate to https://56-228-34-22.sslip.io/mealy/ in Chrome/Firefox',
    'Show the full Mealy interface — padlock icon in address bar confirming HTTPS, '
    'the three-panel layout, camera feed active (or permission prompt), system status indicators'
)
add_screenshot_item(2,
    'Camera permission dialog',
    'Open Mealy in a fresh browser or incognito tab',
    'The browser\'s "Allow camera access?" dialog — proves getUserMedia() works over HTTPS'
)
add_screenshot_item(3,
    'Mealy inside Kowsite (full-screen panel)',
    'Open https://56-228-34-22.sslip.io → click MEALY card in sidebar',
    'Kowsite dashboard with Mealy open as a full-screen overlay — shows the integration '
    'and the Clients section with the MEALY node card'
)
add_screenshot_item(4,
    'EC2 Security Group rules (AWS Console)',
    'AWS Console → EC2 → Security Groups → launch-wizard-2 → Inbound rules',
    'Show ports 22 (SSH), 80 (HTTP), 443 (HTTPS) open — demonstrates cloud security configuration'
)

add_heading('B. AI in Action', 2)

add_screenshot_item(5,
    'Successful food detection',
    'Point camera at food → click Scan Frame',
    'Detection cards appearing in the right panel with food name, confidence %, and kcal — '
    'the AI is working end-to-end'
)
add_screenshot_item(6,
    'Multiple detections with calorie total',
    'Put several foods in frame → scan',
    'Multiple detection cards and the Total kcal counter at the bottom right'
)
add_screenshot_item(7,
    'Auto-scan mode active',
    'Click "Auto" button',
    'The Auto button highlighted green, scan overlay pulsing on the camera feed — '
    'real-time continuous scanning'
)
add_screenshot_item(8,
    'Model switcher — fruit mode',
    'Click "Fruit + Veg" model button, point at fruit, scan',
    'Fruit mode active (button highlighted), detection card showing a fruit class at 96%+ confidence'
)

add_heading('C. Backend & Infrastructure Evidence', 2)

add_screenshot_item(9,
    'systemd services status (terminal)',
    'Run: systemctl --user status mealy-api mealy-tunnel',
    'Both services showing "active (running)" — proves auto-start and tunnel are working'
)
add_screenshot_item(10,
    'mealy-tunnel SSH process (terminal)',
    'Run: ss -tlnp | grep 5001 (on EC2 via SSH)',
    'Port 5001 listening on EC2 via the reverse tunnel — confirms the tunnel is established'
)
add_screenshot_item(11,
    '/health endpoint JSON (browser or curl)',
    'Navigate to https://56-228-34-22.sslip.io/mealy/health in browser',
    'Raw JSON showing {"status":"ok","food_model":"loaded",...} — API health check passing'
)
add_screenshot_item(12,
    'TLS certificate details',
    'Click the padlock in browser address bar → Connection is secure → Certificate',
    'Certificate details showing: Issued to 56-228-34-22.sslip.io, Issued by Let\'s Encrypt, '
    'valid dates — demonstrates HTTPS setup'
)

add_heading('D. Training & Model Evidence', 2)

add_screenshot_item(13,
    'Training log output',
    'Open file train_optimized.log in the project folder, or show Colab training output',
    'Epoch numbers, loss values, accuracy progression — shows the two-phase training'
)
add_screenshot_item(14,
    'Model file in project folder',
    'Open file manager or terminal: ls -lh models/ keggle/',
    'The .keras model files with their sizes — proof of completed training'
)
add_screenshot_item(15,
    'Project folder structure (terminal)',
    'Run: ls /home/K/Storage/Projects/Mealy/ in terminal',
    'All project files — app.py, model.py, train_optimized.py, load_data.py, ui/, food-101/, etc.'
)

add_heading('E. Git History', 2)

add_screenshot_item(16,
    'Git commit history',
    'Run: git log --oneline in the Mealy project folder',
    'All commits showing clean history under your name — project versioning evidence'
)

add_divider()
add_para(
    'Tip: Use a screen recording tool (e.g. OBS, Loom, or built-in screen recorder) '
    'to capture the camera scanning food in real time — a video clip is more compelling '
    'than a screenshot for demonstrating the live AI.',
    italic=True, color=MID
)

# ══════════════════════════════════════════════════════════════════════════════
#  CLOSING
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading('Summary — What You Built', 1)

add_para(
    'Mealy is a full-stack AI application that combines machine learning, web development, '
    'Linux system administration, and cloud infrastructure. Here is what you can claim to have '
    'built and understood:'
)

summary_items = [
    ('AI / ML', 'Trained a convolutional neural network using transfer learning on 75,750 labelled food images. Applied data augmentation to improve generalisation. Used a two-phase training strategy (frozen base + fine-tuning).'),
    ('Backend', 'Built a REST API in Python/Flask that serves the trained model. Implemented lazy model loading, multi-class endpoints, and calorie estimation.'),
    ('Frontend', 'Built a single-page web application with live camera access, real-time scan results, model switching, and session statistics — all in vanilla HTML/CSS/JS.'),
    ('Linux', 'Configured systemd user services for auto-start. Used SSH reverse port forwarding to expose a local server through a cloud machine.'),
    ('Cloud', 'Deployed to Amazon EC2. Configured Nginx as a reverse proxy for two backend services (Flask via tunnel + Node.js Kowsite). Managed AWS Security Groups.'),
    ('Security', 'Obtained a TLS certificate from Let\'s Encrypt using sslip.io as a domain alias. Set up automatic HTTPS with HTTP redirect. Understood why camera access requires secure contexts.'),
    ('Integration', 'Embedded Mealy inside an existing web dashboard (Kowsite) using an iframe, with proper camera delegation and lazy loading.'),
]
for area, desc in summary_items:
    add_bullet(desc, bold_prefix=f'{area}: ')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('mealy · https://56-228-34-22.sslip.io/mealy/')
r.font.size = Pt(11)
r.font.italic = True
r.font.color.rgb = ORANGE

# ── save ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
print(f"Pages: ~{len(doc.paragraphs) // 25 + 1} (estimate)")
